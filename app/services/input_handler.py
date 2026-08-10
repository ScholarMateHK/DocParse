"""
输入模块 - 文件处理服务
支持格式: 纯文本, .txt, .pdf, .doc/.docx
支持分段模式: rule (规则分段), semantic (语义分段)
"""
import os
import re
import shutil
import subprocess
import tempfile
from typing import List, Optional, Tuple, Literal
from pathlib import Path

from docx import Document
from pypdf import PdfReader


# 分段模式类型
SegmentMode = Literal["rule", "semantic"]


class InputHandler:
    """文件输入处理器"""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.doc', '.docx'}
    
    def __init__(self, segment_mode: SegmentMode = "rule"):
        """
        初始化文件处理器
        
        Args:
            segment_mode: 分段模式
                - "rule": 规则分段（基于换行符和标点，速度快）
                - "semantic": 语义分段（基于LLM，效果好但较慢）
        """
        self.ocr_service = None  # 延迟加载OCR服务
        self.segment_mode = segment_mode
        self._semantic_segmenter = None  # 延迟加载语义分段器
    
    def process_input(
        self,
        text: Optional[str] = None,
        file_content: Optional[bytes] = None,
        filename: Optional[str] = None
    ) -> List[str]:
        """
        处理输入，返回分段后的文本列表
        
        Args:
            text: 纯文本输入
            file_content: 文件二进制内容
            filename: 文件名（用于判断文件类型）
            
        Returns:
            List[str]: 分段后的文本列表
        """
        if text:
            return self._process_text(text)
        
        if file_content and filename:
            return self._process_file(file_content, filename)
        
        raise ValueError("必须提供text或file参数")
    
    @property
    def semantic_segmenter(self):
        """延迟加载语义分段器"""
        if self._semantic_segmenter is None:
            from .semantic_segmenter import SemanticSegmenter
            self._semantic_segmenter = SemanticSegmenter()
        return self._semantic_segmenter
    
    def set_segment_mode(self, mode: SegmentMode):
        """设置分段模式"""
        if mode not in ("rule", "semantic"):
            raise ValueError(f"不支持的分段模式: {mode}，支持 'rule' 或 'semantic'")
        self.segment_mode = mode
    
    def _process_text(self, text: str) -> List[str]:
        """处理纯文本输入"""
        # 清理文本
        text = self._clean_text(text)
        
        # 根据模式选择分段方法
        if self.segment_mode == "semantic":
            print("使用语义分段模式（基于LLM）...")
            segments = self.semantic_segmenter.segment(text)
        else:
            # 默认使用规则分段
            segments = self._segment_text(text)
        
        return segments
    
    def _process_file(self, content: bytes, filename: str) -> List[str]:
        """根据文件类型处理文件"""
        ext = Path(filename).suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")
        
        if ext == '.txt':
            return self._process_txt(content)
        elif ext == '.pdf':
            return self._process_pdf(content)
        elif ext in {'.doc', '.docx'}:
            return self._process_docx(content, ext)
        
        raise ValueError(f"未知的文件格式: {ext}")
    
    def _process_txt(self, content: bytes) -> List[str]:
        """处理TXT文件"""
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        text = None
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            raise ValueError("无法解码文件内容")
        
        return self._process_text(text)
    
    def _process_pdf(self, content: bytes) -> List[str]:
        """
        处理PDF文件
        优先使用阿里云文档智能处理复杂排版；服务未配置或调用失败时，
        回退到 pypdf 的本地文本提取。
        """
        print("PDF文件处理：优先使用阿里云文档智能")
        
        try:
            return self._process_pdf_with_ocr(content)
        except Exception as e:
            print(f"云端文档解析不可用 ({e})，回退到本地文本提取")
            return self._process_pdf_fallback(content)
    
    def _process_pdf_fallback(self, content: bytes) -> List[str]:
        """
        PDF文本提取备选方案（当云端文档解析不可用时）。
        使用 pypdf 直接提取文本。
        """
        import io
        
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            full_text = '\n'.join(text_parts)
            
            if len(full_text.strip()) < 50:
                raise ValueError(
                    "PDF文件提取文本过少，可能是扫描版PDF。"
                    "请配置阿里云文档智能凭据后重试"
                )
            
            print("使用本地文本提取，复杂排版可能导致文字顺序错乱")
            return self._process_text(full_text)
            
        except Exception as e:
            raise ValueError(f"PDF文件处理失败: {e}")
    
    def _process_pdf_with_ocr(self, content: bytes) -> List[str]:
        """使用OCR处理PDF"""
        from .ocr_service import OCRService
        
        if self.ocr_service is None:
            self.ocr_service = OCRService()
        
        text = self.ocr_service.ocr_pdf(content)
        return self._process_text(text)
    
    def _process_docx(self, content: bytes, ext: str) -> List[str]:
        """
        处理Word文档
        优先使用阿里云文档智能处理复杂排版，失败时回退本地提取。
        支持 .doc 和 .docx 格式
        """
        if ext == '.doc':
            return self._process_doc_with_ocr(content)
        
        print("DOCX文件处理：优先使用阿里云文档智能")
        
        try:
            return self._process_docx_with_ocr(content)
        except Exception as e:
            print(f"云端文档解析不可用 ({e})，回退到本地文本提取")
            return self._process_docx_fallback(content)
    
    def _process_docx_fallback(self, content: bytes) -> List[str]:
        """
        DOCX文本提取备选方案（当OCR不可用时）
        使用python-docx直接提取文本
        """
        import io
        
        try:
            doc_file = io.BytesIO(content)
            document = Document(doc_file)
            
            text_parts = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 同时提取表格中的文本
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_parts)
            
            if len(full_text.strip()) < 50:
                raise ValueError(
                    "DOCX文件提取文本过少。"
                    "请配置阿里云文档智能凭据后重试"
                )
            
            print("使用本地文本提取，复杂排版可能导致文字顺序错乱")
            return self._process_text(full_text)
            
        except Exception as e:
            raise ValueError(f"DOCX文件处理失败: {e}")
    
    def _process_doc_with_ocr(self, content: bytes) -> List[str]:
        """
        处理 .doc 格式文件
        优先使用阿里云文档智能 API（原生支持 .doc），失败时回退到本地工具
        """
        from .ocr_service import OCRService

        if self.ocr_service is None:
            self.ocr_service = OCRService()

        try:
            text = self.ocr_service.ocr_doc(content)
            return self._process_text(text)
        except Exception as e:
            print(f"云端文档解析失败 ({e})，尝试本地工具提取")

        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
            tmp_file.write(content)
            tmp_path = tmp_file.name

        try:
            text = self._extract_doc_text_fallback(content, tmp_path)
            if text and len(text.strip()) >= 50:
                return self._process_text(text)

            raise ValueError(
                ".doc 文件解析失败，请将文件转换为 .docx 格式后重试"
            )
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    def _extract_doc_text_fallback(self, content: bytes, tmp_path: str) -> Optional[str]:
        """
        .doc文件文本提取备选方案（当OCR不可用时）
        尝试使用 antiword/catdoc 等工具
        """
        text = None
        
        # 方法1: 尝试使用 antiword
        if shutil.which('antiword'):
            try:
                result = subprocess.run(
                    ['antiword', '-w', '0', tmp_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    text = result.stdout
                    print("备选方案：使用 antiword 提取文本")
                    return text
            except Exception as e:
                print(f"antiword 处理失败: {e}")
        
        # 方法2: 尝试使用 catdoc
        if shutil.which('catdoc'):
            try:
                result = subprocess.run(
                    ['catdoc', '-w', tmp_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    text = result.stdout
                    print("备选方案：使用 catdoc 提取文本")
                    return text
            except Exception as e:
                print(f"catdoc 处理失败: {e}")
        
        # 方法3: 尝试从二进制中提取
        try:
            text = self._extract_text_from_binary(content)
            if text and len(text.strip()) > 100:
                print("备选方案：使用二进制提取")
                return text
        except Exception as e:
            print(f"二进制提取失败: {e}")
        
        return None
    
    def _extract_text_from_binary(self, content: bytes) -> Optional[str]:
        """
        从二进制文件中提取可读文本（作为最后的备选方案）
        """
        import re
        
        # 尝试多种编码
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'gbk']:
            try:
                decoded = content.decode(encoding, errors='ignore')
                # 提取可打印字符
                text_parts = re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffefa-zA-Z0-9\s\.\,\;\:\!\?\(\)\[\]\{\}\"\']+', decoded)
                text = ' '.join(text_parts)
                # 清理多余空白
                text = re.sub(r'\s+', ' ', text).strip()
                if len(text) > 100:
                    return text
            except:
                continue
        return None
    
    def _process_docx_with_ocr(self, content: bytes) -> List[str]:
        """优先使用阿里云文档智能处理 Word 文档。"""
        from .ocr_service import OCRService
        
        if self.ocr_service is None:
            self.ocr_service = OCRService()
        
        text = self.ocr_service.ocr_docx(content)
        return self._process_text(text)
    
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 移除多余空白字符
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        # 移除连续的空白行（保留最多一个）
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 移除行首行尾空白
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        return text.strip()
    
    # ==================== 语义聚合规则 ====================
    
    def _apply_colon_aggregation(self, lines: List[str]) -> List[str]:
        """
        规则 A (冒号换行聚合)：
        识别行尾的冒号（：或:）。冒号后的换行内容通常是该条款的补充说明，
        必须合并入同一段。
        """
        if not lines:
            return lines
        
        result = []
        i = 0
        
        while i < len(lines):
            current_line = lines[i]
            
            # 检查当前行是否以冒号结尾
            if current_line.rstrip().endswith(('：', ':')):
                # 收集冒号后的所有相关内容
                aggregated = [current_line]
                i += 1
                
                # 继续收集后续行，直到遇到新的段落开始标志
                while i < len(lines):
                    next_line = lines[i]
                    
                    # 如果下一行是空行，停止聚合
                    if not next_line.strip():
                        break
                    
                    # 如果下一行是新的独立段落开始（非列表项、非缩进补充）
                    # 检查是否是新段落的开始（有明显的段落标题特征）
                    if self._is_new_paragraph_start(next_line) and not self._is_list_item(next_line):
                        break
                    
                    aggregated.append(next_line)
                    i += 1
                
                result.append('\n'.join(aggregated))
            else:
                result.append(current_line)
                i += 1
        
        return result
    
    def _is_list_item(self, line: str) -> bool:
        """
        检查是否是列表项
        支持格式：1. 2. 3.、(1) (2) (3)、① ② ③、第一条、第二条 等
        """
        line = line.strip()
        if not line:
            return False
        
        # 数字+点号：1. 2. 3. 或 1、2、3、
        if re.match(r'^\d+[\.、]\s*', line):
            return True
        
        # 括号数字：(1) (2) (3) 或 （1）（2）（3）
        if re.match(r'^[\(（]\d+[\)）]\s*', line):
            return True
        
        # 圆圈数字：① ② ③
        if re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]\s*', line):
            return True
        
        # 中文序号：第一条、第二条 或 第一、第二
        if re.match(r'^第[一二三四五六七八九十百]+[条款项节章部分]?\s*', line):
            return True
        
        # 字母序号：a. b. c. 或 A. B. C. 或 a) b) c)
        if re.match(r'^[a-zA-Z][\.、\)）]\s*', line):
            return True
        
        # 中文数字序号：一、二、三、或 一. 二. 三.
        if re.match(r'^[一二三四五六七八九十]+[、\.]\s*', line):
            return True
        
        # 破折号或bullet：- • ·
        if re.match(r'^[-•·]\s+', line):
            return True
        
        return False
    
    def _get_list_item_level(self, line: str) -> Tuple[str, int]:
        """
        获取列表项的类型和级别
        返回 (类型标识, 序号数值)
        """
        line = line.strip()
        
        # 数字+点号：1. 2. 3.
        match = re.match(r'^(\d+)[\.、]\s*', line)
        if match:
            return ('num_dot', int(match.group(1)))
        
        # 括号数字：(1) (2) (3)
        match = re.match(r'^[\(（](\d+)[\)）]\s*', line)
        if match:
            return ('paren_num', int(match.group(1)))
        
        # 圆圈数字
        circle_nums = '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳'
        for i, c in enumerate(circle_nums, 1):
            if line.startswith(c):
                return ('circle', i)
        
        # 第X条
        match = re.match(r'^第([一二三四五六七八九十百]+)[条款项节章部分]?\s*', line)
        if match:
            return ('di_tiao', self._chinese_num_to_int(match.group(1)))
        
        # 中文数字序号
        match = re.match(r'^([一二三四五六七八九十]+)[、\.]\s*', line)
        if match:
            return ('cn_num', self._chinese_num_to_int(match.group(1)))
        
        return ('other', 0)
    
    def _chinese_num_to_int(self, cn_num: str) -> int:
        """将中文数字转换为阿拉伯数字"""
        cn_num_map = {
            '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
            '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
            '百': 100
        }
        
        if len(cn_num) == 1:
            return cn_num_map.get(cn_num, 0)
        
        result = 0
        temp = 0
        for char in cn_num:
            if char == '十':
                if temp == 0:
                    temp = 1
                result += temp * 10
                temp = 0
            elif char == '百':
                if temp == 0:
                    temp = 1
                result += temp * 100
                temp = 0
            else:
                temp = cn_num_map.get(char, 0)
        
        result += temp
        return result if result > 0 else 1
    
    def _apply_list_aggregation(self, lines: List[str]) -> List[str]:
        """
        规则 B (列表项汇聚)：
        识别 1. 2. 3.、(1) (2) (3) 或 第一条、第二条 等引导词。
        禁止在中途切断，必须将一整组列表项聚合为一个语义分段（Segment）。
        """
        if not lines:
            return lines
        
        result = []
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 检查是否是列表项的开始
            if self._is_list_item(line):
                list_type, _ = self._get_list_item_level(line)
                
                # 收集整组列表
                list_items = [line]
                i += 1
                
                while i < len(lines):
                    next_line = lines[i]
                    
                    # 空行可能是列表间的分隔，但如果后面还有同类型列表项，继续聚合
                    if not next_line.strip():
                        # 向前看，检查是否还有同类型的列表项
                        lookahead = i + 1
                        while lookahead < len(lines) and not lines[lookahead].strip():
                            lookahead += 1
                        
                        if lookahead < len(lines) and self._is_list_item(lines[lookahead]):
                            next_type, _ = self._get_list_item_level(lines[lookahead])
                            if next_type == list_type:
                                # 保留空行，继续聚合
                                list_items.append(next_line)
                                i += 1
                                continue
                        
                        # 否则，列表结束
                        break
                    
                    # 检查是否是同类型的列表项
                    if self._is_list_item(next_line):
                        next_type, _ = self._get_list_item_level(next_line)
                        if next_type == list_type:
                            list_items.append(next_line)
                            i += 1
                            continue
                        else:
                            # 不同类型的列表，可能是嵌套列表，也聚合
                            list_items.append(next_line)
                            i += 1
                            continue
                    else:
                        # 非列表项，可能是列表项的延续内容（跨行）
                        # 如果这行不是新段落开始，则视为当前列表项的延续
                        if not self._is_new_paragraph_start(next_line):
                            list_items.append(next_line)
                            i += 1
                            continue
                        else:
                            break
                
                result.append('\n'.join(list_items))
            else:
                result.append(line)
                i += 1
        
        return result
    
    def _is_new_paragraph_start(self, line: str) -> bool:
        """
        判断是否是新段落的开始
        """
        line = line.strip()
        if not line:
            return False
        
        # 以数字序号开头通常不是新段落（是列表项）
        if self._is_list_item(line):
            return False
        
        # 以中文标题特征开头
        # 如：一、xxx  第一章 xxx  (一) xxx
        if re.match(r'^[一二三四五六七八九十]+[、\.]\s*\S', line):
            return True
        if re.match(r'^第[一二三四五六七八九十百]+[章节部分条款]\s*', line):
            return True
        if re.match(r'^[\(（][一二三四五六七八九十]+[\)）]\s*\S', line):
            return True
        
        # 以大写字母或数字标题开头
        if re.match(r'^\d+\.\d+\s+', line):  # 如 1.1 xxx
            return True
        
        return False
    
    def _apply_cross_page_compensation(self, lines: List[str]) -> List[str]:
        """
        规则 C (跨页逻辑补偿)：
        跨页处若当前句未结束（无句号/感叹号等终止符），
        需强制与下一页首行拼接，确保语意连贯。
        
        注：由于处理的是已提取的文本，我们通过检测不完整句子来处理
        """
        if not lines:
            return lines
        
        # 句子终止符
        sentence_terminators = ('。', '！', '？', '；', '.', '!', '?', '…', '"', '"', '）', ')')
        # 可能表示跨页的标记（页码等，用于识别可能的页边界）
        page_markers = re.compile(r'^[-—]\s*\d+\s*[-—]$|^\d+$|^第\s*\d+\s*页$|^Page\s*\d+$', re.IGNORECASE)
        
        result = []
        i = 0
        
        while i < len(lines):
            current_line = lines[i]
            
            # 跳过页码标记
            if page_markers.match(current_line.strip()):
                i += 1
                continue
            
            # 检查当前行是否以终止符结尾
            current_stripped = current_line.rstrip()
            
            if current_stripped and not current_stripped.endswith(sentence_terminators):
                # 句子未结束，尝试与后续行拼接
                combined = [current_line]
                i += 1
                
                while i < len(lines):
                    next_line = lines[i]
                    
                    # 跳过页码
                    if page_markers.match(next_line.strip()):
                        i += 1
                        continue
                    
                    # 如果下一行是空行，停止拼接
                    if not next_line.strip():
                        break
                    
                    # 如果下一行是新段落开始或列表项，停止拼接
                    if self._is_new_paragraph_start(next_line) or self._is_list_item(next_line):
                        break
                    
                    combined.append(next_line)
                    i += 1
                    
                    # 检查拼接后是否完成句子
                    if next_line.rstrip().endswith(sentence_terminators):
                        break
                
                # 拼接时去掉换行，保持连贯
                result.append(' '.join(line.strip() for line in combined))
            else:
                result.append(current_line)
                i += 1
        
        return result
    
    def _apply_semantic_aggregation(self, text: str) -> str:
        """
        应用所有语义聚合规则
        """
        # 按行分割
        lines = text.split('\n')
        
        # 规则 C: 跨页逻辑补偿（先处理，确保句子完整）
        lines = self._apply_cross_page_compensation(lines)
        
        # 规则 A: 冒号换行聚合
        lines = self._apply_colon_aggregation(lines)
        
        # 规则 B: 列表项汇聚
        lines = self._apply_list_aggregation(lines)
        
        return '\n'.join(lines)
    
    # ==================== 分段逻辑 ====================
    
    def _segment_text(self, text: str) -> List[str]:
        """
        语义分段
        基于换行符、标点符号和语义完整性进行分段
        
        处理流程：
        1. 应用语义聚合规则（冒号聚合、列表聚合、跨页补偿）
        2. 按双换行分段
        3. 处理过长段落
        4. 合并过短段落
        """
        # 应用语义聚合规则
        text = self._apply_semantic_aggregation(text)
        
        segments = []
        
        # 按双换行分段（明确的段落分隔）
        paragraphs = re.split(r'\n\s*\n', text)
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 如果段落太长，进一步分割（但要保护列表项完整性）
            if len(para) > 3000:
                sub_segments = self._split_long_paragraph_smart(para)
                segments.extend(sub_segments)
            else:
                segments.append(para)
        
        # 合并过短的段落
        segments = self._merge_short_segments(segments)
        
        return segments
    
    def _split_long_paragraph_smart(self, text: str, max_length: int = 3000) -> List[str]:
        """
        智能分割过长的段落，同时保护列表项完整性
        """
        # 如果整个段落是列表，尝试按列表项分组
        lines = text.split('\n')
        
        # 检查是否主要是列表内容
        list_lines = [l for l in lines if self._is_list_item(l.strip())]
        if len(list_lines) > len(lines) * 0.3:  # 超过30%是列表项
            return self._split_list_paragraph(text, max_length)
        
        # 否则按句子分割
        return self._split_long_paragraph(text, max_length)
    
    def _split_list_paragraph(self, text: str, max_length: int = 3000) -> List[str]:
        """
        按列表项分割段落，确保列表项不被中途切断
        """
        lines = text.split('\n')
        segments = []
        current_segment = []
        current_length = 0
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 如果是列表项开始
            if self._is_list_item(line.strip()):
                # 收集这个列表项的完整内容（包括可能的续行）
                item_lines = [line]
                item_length = len(line)
                i += 1
                
                while i < len(lines):
                    next_line = lines[i]
                    # 如果是新的列表项或空行，停止
                    if not next_line.strip() or self._is_list_item(next_line.strip()):
                        break
                    item_lines.append(next_line)
                    item_length += len(next_line) + 1
                    i += 1
                
                # 检查是否需要开始新段落
                if current_length + item_length > max_length and current_segment:
                    segments.append('\n'.join(current_segment))
                    current_segment = []
                    current_length = 0
                
                current_segment.extend(item_lines)
                current_length += item_length
            else:
                # 非列表项，直接添加
                if current_length + len(line) > max_length and current_segment:
                    segments.append('\n'.join(current_segment))
                    current_segment = []
                    current_length = 0
                
                current_segment.append(line)
                current_length += len(line) + 1
                i += 1
        
        if current_segment:
            segments.append('\n'.join(current_segment))
        
        return segments if segments else [text]
    
    def _split_long_paragraph(self, text: str, max_length: int = 2000) -> List[str]:
        """分割过长的段落"""
        segments = []
        
        # 按句子分割
        sentences = re.split(r'([。！？；\n])', text)
        
        current_segment = ""
        for i in range(0, len(sentences), 2):
            sentence = sentences[i]
            # 加上标点符号
            if i + 1 < len(sentences):
                sentence += sentences[i + 1]
            
            if len(current_segment) + len(sentence) > max_length:
                if current_segment:
                    segments.append(current_segment.strip())
                current_segment = sentence
            else:
                current_segment += sentence
        
        if current_segment.strip():
            segments.append(current_segment.strip())
        
        return segments if segments else [text]
    
    def _merge_short_segments(self, segments: List[str], min_length: int = 50) -> List[str]:
        """合并过短的段落"""
        if not segments:
            return segments
        
        merged = []
        current = ""
        
        for segment in segments:
            if len(segment) < min_length:
                current = (current + "\n" + segment).strip() if current else segment
            else:
                if current:
                    # 如果当前累积的短段落加上这个段落不太长，合并
                    if len(current) + len(segment) < 3000:
                        current = current + "\n" + segment
                    else:
                        merged.append(current)
                        current = segment
                else:
                    current = segment
        
        if current:
            merged.append(current)
        
        return merged if merged else segments
