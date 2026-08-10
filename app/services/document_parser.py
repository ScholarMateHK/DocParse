"""
文档解析器 - 整合输入处理和LLM分类

支持两种处理模式：
1. 分离模式(separate): 分段和分类分开调用LLM（兼容旧版）
2. 统一模式(unified): 一次LLM调用同时完成分段和分类（推荐，效率高）

V1版本功能：
- 13类标签体系（包含标题和摘要）
- 标题识别（可为空）
- 摘要识别与自动生成（若原文无摘要则生成500字摘要）
- 可选的二次切分功能（chunk_size + overlap）
"""
import io
import re
import uuid
from typing import List, Optional, Literal, Tuple

from pypdf import PdfReader

from ..models.schemas import Segment, Chunk, ParseResponse
from .input_handler import InputHandler, SegmentMode
from .llm_classifier import LLMClassifier
from ..config import settings


# 处理模式类型
ProcessMode = Literal["separate", "unified"]
# 切分方式类型
ChunkMethod = Literal["sliding", "semantic"]


class DocumentParser:
    """文档解析器主类"""
    
    def __init__(self, segment_mode: SegmentMode = None, process_mode: ProcessMode = None):
        """
        初始化解析器
        
        Args:
            segment_mode: 分段模式（仅在 separate 模式下有效）
                - "rule": 规则分段（基于换行符和标点，速度快）
                - "semantic": 语义分段（基于LLM，效果好但较慢）
                - None: 使用配置文件中的默认值
            process_mode: 处理模式
                - "unified": 统一模式，一次LLM调用完成分段和分类
                - "separate": 分离模式，先分段再批量分类（推荐，效率高）
                - None: 使用配置文件中的默认值
        """
        if segment_mode is None:
            segment_mode = settings.DEFAULT_SEGMENT_MODE
        if process_mode is None:
            process_mode = getattr(settings, 'DEFAULT_PROCESS_MODE', 'separate')
        
        self.input_handler = InputHandler(segment_mode=segment_mode)
        self.classifier = LLMClassifier(batch_mode=getattr(settings, 'DEFAULT_BATCH_MODE', 'batch'))
        self.process_mode = process_mode
        self._unified_processor = None  # 延迟加载
        self._text_chunker = None  # 延迟加载
    
    @property
    def unified_processor(self):
        """延迟加载统一处理器"""
        if self._unified_processor is None:
            from .unified_processor import UnifiedProcessor
            self._unified_processor = UnifiedProcessor()
        return self._unified_processor
    
    @property
    def text_chunker(self):
        """延迟加载文本切分器"""
        if self._text_chunker is None:
            from .text_chunker import TextChunker
            self._text_chunker = TextChunker()
        return self._text_chunker
    
    def set_segment_mode(self, mode: SegmentMode):
        """
        设置分段模式（仅在 separate 模式下有效）
        
        Args:
            mode: "rule" 或 "semantic"
        """
        self.input_handler.set_segment_mode(mode)
    
    def set_process_mode(self, mode: ProcessMode):
        """
        设置处理模式
        
        Args:
            mode: "unified" 或 "separate"
        """
        if mode not in ("unified", "separate"):
            raise ValueError(f"不支持的处理模式: {mode}，支持 'unified' 或 'separate'")
        self.process_mode = mode
    
    def parse(
        self,
        text: Optional[str] = None,
        file_content: Optional[bytes] = None,
        filename: Optional[str] = None,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None,
        chunk_method: ChunkMethod = "sliding"
    ) -> ParseResponse:
        """
        解析文档
        
        Args:
            text: 纯文本输入（可选）
            file_content: 文件二进制内容（可选）
            filename: 文件名（可选）
            chunk_size: 二次切分的块大小（可选，若传入则对段落进行二次切分）
            overlap: 切分重叠区间（可选，仅在chunk_size传入时生效）
            chunk_method: 切分方式 - "sliding"(滑动窗口) 或 "semantic"(语义切分)
            
        Returns:
            ParseResponse: 解析结果
        """
        doc_id = str(uuid.uuid4())
        
        try:
            # 根据处理模式选择不同的处理流程
            if self.process_mode == "unified":
                response = self._parse_unified(doc_id, text, file_content, filename)
            else:
                response = self._parse_separate(doc_id, text, file_content, filename)
            
            # 如果解析失败，直接返回
            if response.status != "success":
                return response
            
            # 提取标题和摘要
            title, abstract, remaining_segments = self._extract_title_and_abstract(
                response.segments,
                text or self._extract_text(file_content, filename) if file_content else None
            )
            
            # 更新响应
            response.title = title
            response.abstract = abstract
            response.segments = remaining_segments
            
            # 如果传入了chunk_size，进行二次切分
            if chunk_size is not None and chunk_size > 0:
                response = self._apply_chunking(
                    response, 
                    chunk_size, 
                    overlap or 0, 
                    chunk_method
                )
            
            return response
            
        except Exception as e:
            print(f"解析失败: {e}")
            import traceback
            traceback.print_exc()
            return ParseResponse(
                status="error",
                doc_id=doc_id,
                segments=[],
                error_message=str(e)
            )
    
    def _extract_title_and_abstract(
        self,
        segments: List[Segment],
        full_text: Optional[str] = None
    ) -> tuple:
        """
        从段落中提取标题和摘要
        
        Args:
            segments: 段落列表
            full_text: 完整文档文本（用于生成摘要）
            
        Returns:
            tuple: (title, abstract, remaining_segments)
        """
        title = None
        abstract = None
        remaining_segments = []
        consumed_segments = []

        for segment in segments:
            if segment.tag == "标题":
                if title is None:
                    title = segment.content.strip()
                    consumed_segments.append(segment)
                else:
                    segment.tag = "其他"
                    remaining_segments.append(segment)
            elif segment.tag == "摘要":
                if abstract is None:
                    abstract = segment.content.strip()
                else:
                    abstract = abstract + "\n" + segment.content.strip()
                consumed_segments.append(segment)
            else:
                remaining_segments.append(segment)

        # 如果没有摘要，调用LLM生成
        if abstract is None or not abstract.strip():
            print("未找到摘要，正在使用LLM生成...")
            content_for_abstract = full_text
            if not content_for_abstract:
                content_for_abstract = "\n".join([s.content for s in remaining_segments])

            if content_for_abstract and content_for_abstract.strip():
                abstract = self.classifier.generate_abstract(content_for_abstract, max_length=500)
                if abstract:
                    print(f"成功生成摘要，长度: {len(abstract)} 字")
                else:
                    print("摘要生成失败")

        # segments 为空时的兜底处理
        if not remaining_segments:
            extracted_texts = set()
            if title:
                extracted_texts.add(title)
            if abstract:
                extracted_texts.add(abstract)

            if full_text and full_text.strip():
                leftover = full_text
                for t in extracted_texts:
                    leftover = leftover.replace(t, "", 1)
                leftover = leftover.strip()
                if leftover:
                    remaining_segments.append(Segment(
                        segment_id=1,
                        content=leftover,
                        tag="其他",
                        confidence=1.0
                    ))

            if not remaining_segments and consumed_segments:
                for seg in consumed_segments:
                    seg.tag = "其他"
                    remaining_segments.append(seg)

        # 重新编号剩余段落
        for i, segment in enumerate(remaining_segments):
            segment.segment_id = i + 1

        return title, abstract, remaining_segments
    
    def _apply_chunking(
        self,
        response: ParseResponse,
        chunk_size: int,
        overlap: int,
        chunk_method: ChunkMethod
    ) -> ParseResponse:
        """
        对段落进行二次切分
        
        Args:
            response: 原始解析响应
            chunk_size: 切分块大小
            overlap: 重叠区间
            chunk_method: 切分方式
            
        Returns:
            ParseResponse: 包含chunks的响应
        """
        print(f"正在进行二次切分 (chunk_size={chunk_size}, overlap={overlap}, method={chunk_method})...")
        
        # 转换segments为dict格式
        segments_dict = [
            {
                'segment_id': s.segment_id,
                'content': s.content,
                'tag': s.tag,
                'confidence': s.confidence,
                'pageIdx': s.pageIdx,
            }
            for s in response.segments
        ]
        
        # 进行切分
        chunked_segments = self.text_chunker.chunk_segments(
            segments=segments_dict,
            chunk_size=chunk_size,
            overlap=overlap,
            method=chunk_method
        )
        
        # 转换回Segment对象
        new_segments = []
        for seg_data in chunked_segments:
            chunks = None
            if seg_data.get('chunks'):
                chunks = [
                    Chunk(
                        chunk_id=c['chunk_id'],
                        segment_id=c['segment_id'],
                        content=c['content'],
                        tag=c['tag'],
                        start_pos=c['start_pos'],
                        end_pos=c['end_pos']
                    )
                    for c in seg_data['chunks']
                ]
            
            new_segments.append(Segment(
                segment_id=seg_data['segment_id'],
                content=seg_data['content'],
                tag=seg_data['tag'],
                confidence=seg_data['confidence'],
                pageIdx=seg_data.get('pageIdx'),
                chunks=chunks,
            ))
        
        response.segments = new_segments
        response.chunk_info = {
            'chunk_size': chunk_size,
            'overlap': overlap,
            'method': chunk_method
        }
        
        # 统计切分结果
        total_chunks = sum(
            len(s.chunks) if s.chunks else 0 
            for s in new_segments
        )
        print(f"二次切分完成，共生成 {total_chunks} 个chunks")
        
        return response
    
    def _parse_unified(
        self,
        doc_id: str,
        text: Optional[str] = None,
        file_content: Optional[bytes] = None,
        filename: Optional[str] = None
    ) -> ParseResponse:
        """
        统一模式解析：一次 LLM 调用完成分段和分类
        """
        print("使用统一模式（unified）：一次LLM调用完成分段和分类...")
        
        per_page_texts = None
        if file_content and filename:
            try:
                raw_text, per_page_texts = self._extract_text_with_pages(file_content, filename)
            except Exception:
                raw_text = self._extract_text(file_content, filename)
        elif text:
            raw_text = text
        else:
            raise ValueError("必须提供text或file参数")
        
        if not raw_text or not raw_text.strip():
            return ParseResponse(
                status="error",
                doc_id=doc_id,
                segments=[],
                error_message="文档内容为空"
            )
        
        # 使用统一处理器一次性完成分段和分类
        print("调用LLM进行统一处理（分段+分类）...")
        results = self.unified_processor.process(raw_text)
        
        if not results:
            return ParseResponse(
                status="error",
                doc_id=doc_id,
                segments=[],
                error_message="处理结果为空"
            )
        
        # 构建响应
        segments = []
        for i, result in enumerate(results):
            segment = Segment(
                segment_id=i + 1,
                content=result.content,
                tag=result.tag,
                confidence=round(result.confidence, 2)
            )
            segments.append(segment)
        
        if per_page_texts:
            self._assign_pages_to_segments(segments, per_page_texts)
        
        print(f"统一模式解析完成，共 {len(segments)} 个段落")
        
        return ParseResponse(
            status="success",
            doc_id=doc_id,
            segments=segments
        )
    
    def _parse_separate(
        self,
        doc_id: str,
        text: Optional[str] = None,
        file_content: Optional[bytes] = None,
        filename: Optional[str] = None
    ) -> ParseResponse:
        """
        分离模式解析：分段和分类分开调用 LLM（兼容旧版）
        """
        print("使用分离模式（separate）：分段和分类分开调用LLM...")
        
        # 步骤1：处理输入，获取分段后的文本
        print("步骤1: 处理输入文件/文本进行分段...")
        per_page_texts = None
        if file_content and filename:
            try:
                raw_text, per_page_texts = self._extract_text_with_pages(file_content, filename)
                segments_text = self.input_handler.process_input(text=raw_text)
            except Exception:
                segments_text = self.input_handler.process_input(
                    file_content=file_content, filename=filename
                )
                per_page_texts = None
        else:
            segments_text = self.input_handler.process_input(text=text)
        
        if not segments_text:
            return ParseResponse(
                status="error",
                doc_id=doc_id,
                segments=[],
                error_message="文档内容为空"
            )
        
        print(f"文档分段完成，共 {len(segments_text)} 个段落")
        
        # 步骤2：使用LLM对每个段落进行分类
        print("步骤2: 使用LLM进行段落分类...")
        classification_results = self.classifier.classify_segments_batch(segments_text)
        
        # 步骤3：构建响应
        segments = []
        for i, (content, (tag, confidence)) in enumerate(
            zip(segments_text, classification_results)
        ):
            segment = Segment(
                segment_id=i + 1,
                content=content,
                tag=tag,
                confidence=round(confidence, 2)
            )
            segments.append(segment)
        
        if per_page_texts:
            self._assign_pages_to_segments(segments, per_page_texts)
        
        print(f"分离模式解析完成，共识别 {len(segments)} 个段落")
        
        return ParseResponse(
            status="success",
            doc_id=doc_id,
            segments=segments
        )
    
    def _extract_text_with_pages(
        self,
        file_content: bytes,
        filename: str
    ) -> Tuple[str, Optional[List[str]]]:
        """
        从文件中提取文本并追踪页码边界。

        Returns:
            (full_text, per_page_texts)
            per_page_texts[i] 对应第 i+1 页的文本。
            当无法获取页码时 per_page_texts 为 None。
        """
        from pathlib import Path

        ext = Path(filename).suffix.lower()

        # 初始化阿里云文档智能服务
        if self.input_handler.ocr_service is None:
            from .ocr_service import OCRService
            try:
                self.input_handler.ocr_service = OCRService()
            except Exception:
                pass

        # 云端文档智能路径：支持 PDF / DOC / DOCX，可获取按页分组的文本
        if self.input_handler.ocr_service and ext in {'.pdf', '.doc', '.docx'}:
            try:
                text, per_page_texts = (
                    self.input_handler.ocr_service.parse_file_with_pages(
                        file_content, filename
                    )
                )
                if text and text.strip():
                    return text, per_page_texts
            except Exception as e:
                print(f"云端文档解析失败 ({e})，回退到本地提取")

        # PDF 备选：pypdf 逐页提取（仍可获取页码）
        if ext == '.pdf':
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            per_page_texts: List[str] = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    per_page_texts.append(page_text)
            if per_page_texts:
                return '\n'.join(per_page_texts), per_page_texts
            raise ValueError("PDF提取文本为空")

        # TXT：无页码概念
        if ext == '.txt':
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    return file_content.decode(encoding), None
                except UnicodeDecodeError:
                    continue
            raise ValueError("无法解码文件内容")

        # DOCX 备选：python-docx（无页码信息）
        if ext == '.docx':
            from docx import Document
            doc_file = io.BytesIO(file_content)
            document = Document(doc_file)
            text_parts = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            return '\n'.join(text_parts), None

        raise ValueError(f"_extract_text_with_pages 不支持: {ext}")

    def _assign_pages_to_segments(
        self,
        segments: List[Segment],
        per_page_texts: List[str]
    ) -> None:
        """根据内容匹配为每个 segment 分配 pageIdx（起始页码）。"""
        if not per_page_texts:
            return
        normalized_pages = [re.sub(r'\s+', '', pt) for pt in per_page_texts]
        for segment in segments:
            segment.pageIdx = self._find_segment_page(
                segment.content, normalized_pages
            )

    @staticmethod
    def _find_segment_page(
        content: str,
        normalized_pages: List[str]
    ) -> Optional[int]:
        """通过子串采样定位 segment 所在的起始页码。"""
        normalized = re.sub(r'\s+', '', content)
        if not normalized:
            return None

        sample_len = min(60, len(normalized))
        sample = normalized[:sample_len]

        for page_idx, page_norm in enumerate(normalized_pages):
            if sample in page_norm:
                return page_idx + 1

        short_len = min(30, len(normalized))
        short_sample = normalized[:short_len]
        for page_idx, page_norm in enumerate(normalized_pages):
            if short_sample in page_norm:
                return page_idx + 1

        return None

    def _extract_text(self, file_content: bytes, filename: str) -> str:
        """
        从文件中提取原始文本（不分段）
        """
        from pathlib import Path
        import io
        
        ext = Path(filename).suffix.lower()
        
        if ext == '.txt':
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("无法解码文件内容")
        
        elif ext == '.pdf':
            # 优先使用云端文档智能，失败时回退到 pypdf
            if self.input_handler.ocr_service is None:
                from .ocr_service import OCRService
                try:
                    self.input_handler.ocr_service = OCRService()
                except Exception:
                    pass
            
            if self.input_handler.ocr_service:
                try:
                    return self.input_handler.ocr_service.ocr_pdf(file_content)
                except Exception as e:
                    print(f"云端文档解析失败 ({e})，回退到 pypdf")

            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return '\n'.join(text_parts)
        
        elif ext in {'.doc', '.docx'}:
            # 优先使用云端文档智能，失败时回退到本地提取
            if self.input_handler.ocr_service is None:
                from .ocr_service import OCRService
                try:
                    self.input_handler.ocr_service = OCRService()
                except Exception:
                    pass
            
            if ext == '.doc':
                segments = self.input_handler._process_doc_with_ocr(file_content)
                return '\n'.join(segments)
            elif self.input_handler.ocr_service:
                try:
                    return self.input_handler.ocr_service.ocr_docx(file_content)
                except Exception as e:
                    print(f"云端文档解析失败 ({e})，回退到 python-docx")

            from docx import Document
            doc_file = io.BytesIO(file_content)
            document = Document(doc_file)
            text_parts = [p.text for p in document.paragraphs if p.text.strip()]
            return '\n'.join(text_parts)
        
        raise ValueError(f"不支持的文件格式: {ext}")
    
    def parse_with_segment_grouping(
        self,
        text: Optional[str] = None,
        file_content: Optional[bytes] = None,
        filename: Optional[str] = None
    ) -> ParseResponse:
        """
        解析文档并按标签分组段落ID
        
        实现需求中的要求：
        "chunk的ID，分chunk的话则需要保证在一个语义段下chunk ID的连续性，
        如项目研究意义Segment1-SegmentN"
        """
        # 首先进行常规解析
        response = self.parse(
            text=text,
            file_content=file_content,
            filename=filename
        )
        
        if response.status != "success" or not response.segments:
            return response
        
        # 按标签重新编号，保证同一标签下的ID连续
        tag_counters = {}
        renamed_segments = []
        
        for segment in response.segments:
            tag = segment.tag
            if tag not in tag_counters:
                tag_counters[tag] = 1
            
            # 创建新的segment_id格式：原始顺序ID + 标签内序号
            segment_id = segment.segment_id
            tag_index = tag_counters[tag]
            tag_counters[tag] += 1
            
            renamed_segments.append(Segment(
                segment_id=segment_id,
                content=segment.content,
                tag=f"{tag}_Segment{tag_index}",
                confidence=segment.confidence,
                pageIdx=segment.pageIdx,
            ))
        
        response.segments = renamed_segments
        return response
